from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables" / "EBOOK_RV_TRINCA_RV21_PREMIUM_V1.docx"
LOGO = ROOT / "public" / "images" / "logorv.jpg"
PRODUCT = ROOT / "public" / "images" / "trinca-rv21-produto.png"
FACE = ROOT / "public" / "images" / "ruria-rosto-premium.png"

BLACK = "090908"
INK = "191713"
MUTED = "6E675C"
GOLD = "D7B65D"
GOLD_DARK = "8C6A20"
PAPER = "FAF7EF"
SOFT = "F3EBDD"
LINE = "D8C79C"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=LINE, size="8"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa=9360):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_run(run, size=None, color=None, bold=None, italic=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_para(doc, text="", style=None, align=None, before=None, after=None, line=1.25):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_run(run)
    if align is not None:
        p.alignment = align
    fmt = p.paragraph_format
    if before is not None:
        fmt.space_before = Pt(before)
    if after is not None:
        fmt.space_after = Pt(after)
    fmt.line_spacing = line
    return p


def add_rich_para(doc, parts, style=None, align=None, before=None, after=None, line=1.25):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    fmt = p.paragraph_format
    if before is not None:
        fmt.space_before = Pt(before)
    if after is not None:
        fmt.space_after = Pt(after)
    fmt.line_spacing = line
    for text, kwargs in parts:
        run = p.add_run(text)
        set_run(run, **kwargs)
    return p


def add_heading(doc, text, level=1):
    style = "Heading {}".format(level)
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if level == 1:
        set_run(run, size=16, color=GOLD_DARK, bold=True)
    elif level == 2:
        set_run(run, size=13, color=INK, bold=True)
    else:
        set_run(run, size=12, color=GOLD_DARK, bold=True)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        set_run(run, size=10.5, color=INK)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        set_run(run, size=10.5, color=INK)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, SOFT)
    set_cell_border(cell, LINE, "8")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title.upper())
    set_run(r, size=9, color=GOLD_DARK, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.line_spacing = 1.2
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run(r2, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_label_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table)
    table.autofit = False
    hdr = table.rows[0].cells
    hdr[0].text = "Conceito"
    hdr[1].text = "Como aplicar no TRINCA RV21"
    for cell in hdr:
        set_cell_shading(cell, BLACK)
        set_cell_border(cell, BLACK, "8")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=9.5, color=GOLD, bold=True)
    for label, detail in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = detail
        for i, cell in enumerate(cells):
            set_cell_border(cell, LINE, "6")
            set_cell_shading(cell, "FFFFFF" if i else PAPER)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.18
                for run in p.runs:
                    set_run(run, size=9.5 if i else 9.2, color=INK, bold=(i == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(5)
    return table


def add_checklist(doc, items):
    table = doc.add_table(rows=0, cols=2)
    set_table_width(table)
    for item in items:
        row = table.add_row().cells
        row[0].text = "□"
        row[1].text = item
        for i, cell in enumerate(row):
            set_cell_border(cell, LINE, "4")
            set_cell_shading(cell, "FFFFFF")
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_run(run, size=11 if i == 0 else 9.5, color=GOLD_DARK if i == 0 else INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(5)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, GOLD_DARK, 18, 10),
        ("Heading 2", 13, INK, 14, 7),
        ("Heading 3", 12, GOLD_DARK, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    return doc


def add_header_footer(doc):
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run("TRINCA RV21 | Ebook RV")
        set_run(r, size=8.5, color=MUTED, bold=True)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.add_run("@ruriavirginio")
        set_run(fr, size=8.5, color=MUTED)


def add_cover(doc):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, BLACK)
    set_cell_border(cell, BLACK, "12")
    for _ in range(2):
        cell.add_paragraph()
    if LOGO.exists():
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(LOGO), width=Inches(1.05))
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("EBOOK RV")
    set_run(r, size=28, color="FFFFFF", bold=True)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(12)
    r2 = p2.add_run("GUIA DE EXECUÇÃO INTELIGENTE")
    set_run(r2, size=16, color=GOLD, bold=True)
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.line_spacing = 1.15
    r3 = p3.add_run("Consciência corporal, intensidade, constância e método para viver melhor os 21 dias do TRINCA RV21.")
    set_run(r3, size=12, color="FFFFFF")
    if PRODUCT.exists():
        p4 = cell.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p4.paragraph_format.space_before = Pt(18)
        run = p4.add_run()
        run.add_picture(str(PRODUCT), width=Inches(1.25))
    p5 = cell.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p5.paragraph_format.space_before = Pt(18)
    r5 = p5.add_run("Ruriá Virgínio | TRINCA RV21")
    set_run(r5, size=11, color=GOLD, bold=True)
    for _ in range(2):
        cell.add_paragraph()
    doc.add_page_break()


def build():
    OUT.parent.mkdir(exist_ok=True)
    doc = setup_document()
    add_header_footer(doc)
    add_cover(doc)

    add_heading(doc, "Antes de começar", 1)
    add_rich_para(
        doc,
        [
            ("Este ebook não foi criado para ser apenas lido. ", {"size": 11, "color": INK, "bold": True}),
            ("Ele foi criado para orientar sua execução dentro do TRINCA RV21, para que você treine com mais consciência, menos improviso e mais compromisso com o processo.", {"size": 11, "color": INK}),
        ],
        after=8,
    )
    add_callout(
        doc,
        "A decisão central",
        "O resultado dos próximos 21 dias não nasce de perfeição. Ele nasce de repetição, presença e correção de rota. Você não precisa fazer tudo perfeito. Você precisa voltar para o plano todos os dias.",
    )
    add_heading(doc, "Como usar este guia", 2)
    add_numbered(
        doc,
        [
            "Leia uma vez antes de iniciar os treinos.",
            "Volte neste material sempre que tiver dúvida sobre método, descanso, intensidade ou execução.",
            "Use os checklists para corrigir a postura mental e prática antes de treinar.",
            "Não misture este processo com várias estratégias externas ao mesmo tempo.",
        ],
    )

    add_heading(doc, "O método RV dentro do TRINCA", 1)
    add_para(
        doc,
        "O TRINCA RV21 foi construído para transformar tentativa em direção. A aluna não precisa adivinhar o que fazer: ela precisa entender a lógica do treino, aplicar com disciplina e repetir o processo até o corpo começar a responder.",
    )
    add_label_table(
        doc,
        [
            ("Direção", "Você sabe o que executar e por que aquela etapa existe."),
            ("Consciência", "Você presta atenção no movimento, na respiração, na postura e no músculo trabalhado."),
            ("Intensidade", "Você treina com esforço real, respeitando técnica e segurança."),
            ("Constância", "Você mantém o plano mesmo quando a motivação oscila."),
        ],
    )

    add_heading(doc, "Princípios que realmente importam", 1)
    add_para(
        doc,
        "Treinar melhor não é apenas fazer mais exercícios. É compreender como o corpo responde aos estímulos e como pequenas decisões repetidas mudam o resultado.",
    )
    add_label_table(
        doc,
        [
            ("Individualidade biológica", "Cada corpo responde de um jeito. Sono, alimentação, rotina, histórico e metabolismo influenciam o ritmo de evolução."),
            ("Sobrecarga", "O corpo precisa ser desafiado aos poucos: mais carga, mais controle, mais repetições ou melhor execução."),
            ("Adaptação", "Quando o treino fica fácil demais, o corpo para de receber motivo suficiente para mudar."),
            ("Continuidade", "O corpo confia no que você repete. Resultado duradouro nasce de frequência, não de empolgação isolada."),
            ("Especificidade", "Cada objetivo pede uma estratégia: emagrecer, melhorar firmeza, recuperar autoestima ou voltar a vestir roupas antigas exige foco prático."),
            ("Recuperação", "Descanso, hidratação e alimentação fazem parte do resultado. Evolução não acontece só durante o treino."),
        ],
    )

    add_heading(doc, "Execução: onde o resultado começa", 1)
    add_para(
        doc,
        "A execução é o ponto onde muitas pessoas perdem resultado sem perceber. O movimento precisa ter controle, amplitude segura e intenção. Fazer rápido demais, roubar movimento ou treinar sem atenção reduz a qualidade do estímulo.",
    )
    add_heading(doc, "Checklist da boa execução", 2)
    add_checklist(
        doc,
        [
            "Sei qual músculo estou tentando trabalhar.",
            "Consigo controlar a descida e a subida do movimento.",
            "Não estou usando impulso para terminar a repetição.",
            "Minha respiração está organizada.",
            "A carga desafia, mas não destrói minha técnica.",
        ],
    )
    add_callout(
        doc,
        "Regra prática",
        "Se a carga faz você perder postura, amplitude e controle, ela deixou de ser ferramenta de evolução e virou distração.",
    )

    add_heading(doc, "Intensidade sem confusão", 1)
    add_para(
        doc,
        "Intensidade não é treinar de qualquer jeito até ficar exausta. Intensidade é entregar esforço com técnica. Em muitos exercícios, você deve terminar a série sentindo que poderia fazer poucas repetições a mais mantendo boa forma.",
    )
    add_label_table(
        doc,
        [
            ("Leve demais", "Você termina a série sem esforço real. O corpo recebe pouco estímulo."),
            ("Ponto ideal", "As últimas repetições exigem foco, mas a técnica continua limpa."),
            ("Pesado demais", "Você perde amplitude, postura, controle ou sente dor inadequada."),
        ],
    )

    add_heading(doc, "Métodos aplicados no treino", 1)
    add_para(
        doc,
        "Os métodos abaixo aparecem para organizar intensidade. Eles não são enfeite: cada um tem uma função. Use com atenção ao treino do dia e às orientações recebidas.",
    )
    add_label_table(
        doc,
        [
            ("Série simples", "Método base: execute as séries do exercício com descanso adequado antes de avançar."),
            ("Bi-set", "Dois exercícios em sequência, com pouco ou nenhum descanso entre eles."),
            ("Tri-set", "Três exercícios em sequência para aumentar densidade e resistência muscular."),
            ("Drop set", "Após chegar perto da falha, reduza a carga e continue o movimento com controle."),
            ("Rest-pause", "Pequena pausa estratégica para conseguir mais repetições com qualidade."),
            ("Falha muscular", "Momento em que você não consegue outra repetição com técnica segura."),
        ],
    )

    add_heading(doc, "Conceitos rápidos para treinar melhor", 1)
    add_label_table(
        doc,
        [
            ("Cadência", "Velocidade do movimento. Controle gera consciência e melhora a qualidade do estímulo."),
            ("Amplitude", "Movimento completo dentro de uma faixa segura. Encurtar sem necessidade reduz eficiência."),
            ("Progressão de carga", "Aumentar estímulos aos poucos, sem sacrificar técnica."),
            ("Tempo de descanso", "Descanso curto demais pode derrubar performance; longo demais pode reduzir ritmo. Siga a orientação do treino."),
            ("Aquecimento", "Preparação articular e muscular para executar melhor e reduzir risco."),
            ("Cardio", "Ferramenta para condicionamento, saúde cardiovascular e apoio ao gasto energético."),
        ],
    )

    add_heading(doc, "Erros que mais travam evolução", 1)
    add_bullets(
        doc,
        [
            "Começar forte demais e abandonar quando a rotina aperta.",
            "Trocar de estratégia antes de dar tempo ao corpo responder.",
            "Treinar sem registrar percepção, carga ou dificuldade.",
            "Fazer dieta perfeita por dois dias e desistir no primeiro deslize.",
            "Comparar seu início com o meio do processo de outra pessoa.",
            "Confundir cansaço emocional com incapacidade.",
        ],
    )

    add_heading(doc, "O pacto dos 21 dias", 1)
    add_para(
        doc,
        "Durante o TRINCA RV21, o compromisso não é com uma versão impossível de você. O compromisso é com uma versão mais presente, mais firme e mais honesta na execução.",
    )
    add_checklist(
        doc,
        [
            "Vou cumprir o treino previsto ou comunicar minha adaptação quando necessário.",
            "Vou seguir a dieta do meu objetivo sem buscar atalhos confusos.",
            "Vou voltar para o plano depois de um erro, sem transformar um deslize em desistência.",
            "Vou respeitar meu processo sem usar comparação como punição.",
            "Vou tratar esses 21 dias como uma decisão, não como uma tentativa solta.",
        ],
    )

    add_heading(doc, "Mensagem final", 1)
    add_rich_para(
        doc,
        [
            ("Resultado é consequência de constância. ", {"size": 12, "color": INK, "bold": True}),
            ("Não existe transformação real sem processo, mas também não existe processo forte sem direção. Você agora tem uma base mais clara para treinar com inteligência, corrigir sua execução e continuar mesmo nos dias em que a motivação não aparecer.", {"size": 11, "color": INK}),
        ],
        after=8,
    )
    add_callout(
        doc,
        "Leve isso com você",
        "Você não entrou no TRINCA RV21 para provar algo para os outros. Você entrou para voltar a agir por você.",
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
